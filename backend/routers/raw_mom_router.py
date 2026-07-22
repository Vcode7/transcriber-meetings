"""
raw_mom_router.py — RAG-based Raw MoM extraction endpoints.

This router is COMPLETELY INDEPENDENT of the existing /mom router.
It does NOT modify or read from the minutes_of_meeting table.
Raw MoM data is stored in the recordings.raw_mom column.

Endpoints
---------
POST /raw-mom/{recording_id}/retrieve        — Retrieve evidence only (no LLM)
POST /raw-mom/{recording_id}/generate        — Accept pre-assembled evidence, call LLM
POST /raw-mom/{recording_id}/generate-full   — Run the full RAG+LLM pipeline (legacy)
GET  /raw-mom/{recording_id}                 — Fetch stored raw_mom JSON
DELETE /raw-mom/{recording_id}               — Clear raw_mom for a recording
GET  /raw-mom/{recording_id}/status          — Pipeline status (embedded flags, etc.)
GET  /raw-mom/{recording_id}/download/docx   — Download raw_mom as DOCX
"""
from __future__ import annotations

import io
import json
import logging
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import text

from database import get_db, dt_to_str, from_json
from routers.auth import get_current_user
from config import settings

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/raw-mom", tags=["raw-mom"])


# ── Pydantic request models ───────────────────────────────────────────────────

class AgendaItem(BaseModel):
    topic: str
    speaker: Optional[str] = None


class GenerateFinalMomRequest(BaseModel):
    char_limit: Optional[int] = None
    use_max_tokens: Optional[bool] = False


class RetrieveRequest(BaseModel):
    k_transcript: int = Field(default=8, ge=0, le=30)
    k_meeting: int = Field(default=4, ge=0, le=30)
    k_global: int = Field(default=2, ge=0, le=30)
    relative_similarity_cutoff: float = Field(default=0.01, ge=0.0, le=1.0)
    char_limit: int = Field(default=15000, ge=1000, le=100000)
    force_reembed: bool = False
    # Agenda-first: pre-parsed agenda avoids a second LLM call at retrieve time
    agenda_items: Optional[List[AgendaItem]] = None
    # Hybrid transcript retrieval params
    timeline_stride_seconds: float = Field(default=60.0, ge=0.0, le=600.0)
    high_confidence_threshold: float = Field(default=0.70, ge=0.0, le=1.0)
    recording_duration: float = Field(default=0.0, ge=0.0)
    retrieve_by_timeline: bool = False
    # Retrieval mode: "agenda_wise" (default) or "chunk_wise"
    retrieval_mode: str = Field(default="agenda_wise")
    # Maximum agendas a single chunk can be assigned to in chunk-wise mode
    max_overlap_chunks: int = Field(default=2, ge=1, le=20)


class AgendaEvidenceItem(BaseModel):
    topic: str
    speaker: Optional[str] = None
    evidence: str  # pre-assembled evidence string from frontend
    manual_context: Optional[str] = None  # per-agenda free-text context added in UI


class GenerateFromEvidenceRequest(BaseModel):
    agendas: List[AgendaEvidenceItem]


# ── Helpers ───────────────────────────────────────────────────────────────────

async def _get_recording_or_404(recording_id: str, user_id: str) -> dict:
    """Fetch recording row or raise 404."""
    async with get_db() as db:
        r = await db.execute(
            text("""
                SELECT id, filename, created_at, duration, speakers_detected,
                       transcript, raw_mom, transcript_embedded, meeting_context_embedded,
                       agenda_summary, context_summary
                FROM recordings
                WHERE id = :id AND user_id = :uid
            """),
            {"id": recording_id, "uid": user_id},
        )
        row = r.mappings().fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Recording not found")
    return dict(row)


# ── GET status ────────────────────────────────────────────────────────────────

@router.get("/{recording_id}/status")
async def get_raw_mom_status(
    recording_id: str,
    current_user: dict = Depends(get_current_user),
):
    """
    Return the RAG pipeline status for a recording:
    - Whether raw_mom has been generated
    - Embedding flags (transcript, meeting context)
    - Embedding model info
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    has_raw_mom = bool(rec.get("raw_mom"))
    raw_mom_summary = None
    if has_raw_mom:
        try:
            data = json.loads(rec["raw_mom"])
            agendas = data.get("meeting", {}).get("agendas", [])
            raw_mom_summary = {
                "agenda_count": len(agendas),
                "total_discussion_entries": sum(
                    len(a.get("discussion", [])) for a in agendas
                ),
            }
        except Exception:
            pass

    return {
        "recording_id": recording_id,
        "has_raw_mom": has_raw_mom,
        "raw_mom_summary": raw_mom_summary,
        "transcript_embedded": bool(rec.get("transcript_embedded")),
        "meeting_context_embedded": bool(rec.get("meeting_context_embedded")),
        "embedding_model": settings.QWEN_EMBEDDING_MODEL_NAME,
        "has_agenda": bool(rec.get("agenda_summary")),
    }


# ── GET raw_mom ───────────────────────────────────────────────────────────────

@router.get("/{recording_id}")
async def get_raw_mom(
    recording_id: str,
    current_user: dict = Depends(get_current_user),
):
    """
    Fetch the stored raw_mom JSON for a recording.

    Returns 404 if raw MoM has not been generated yet.
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    if not rec.get("raw_mom"):
        raise HTTPException(
            status_code=404,
            detail="Raw MoM has not been generated for this recording. "
                   "Use POST /raw-mom/{id}/generate to generate it.",
        )

    try:
        raw_mom_data = json.loads(rec["raw_mom"])
    except Exception:
        raise HTTPException(status_code=500, detail="Raw MoM data is corrupted")

    return raw_mom_data


# ── POST retrieve (retrieval only, no LLM) ───────────────────────────────────

@router.post("/{recording_id}/retrieve")
async def retrieve_evidence_endpoint(
    recording_id: str,
    body: RetrieveRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Retrieve evidence chunks for all agenda items without running any LLM.

    This is the first step of the interactive Raw MoM Lab workflow:
    1. Parse agenda (from DB cache or re-parse)
    2. Embed transcript/meeting context if needed
    3. Run FAISS retrieval per agenda item
    4. Apply relative similarity filtering
    5. Return raw chunks with full metadata

    No LLM call is made — callers can inspect, filter, and reorder chunks
    before calling POST /generate with pre-assembled evidence.
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    transcript = from_json(rec.get("transcript"), [])
    if not transcript:
        raise HTTPException(
            status_code=400,
            detail="No transcript available. Run transcription first.",
        )

    agenda_text: Optional[str] = rec.get("agenda_summary") or None
    if not agenda_text:
        agenda_text = await _load_raw_agenda_text(recording_id, user_id)

    import asyncio
    _loop = asyncio.get_running_loop()

    try:
        result = await _loop.run_in_executor(
            None,
            lambda: _run_retrieve_pipeline(
                recording_id=recording_id,
                user_id=user_id,
                transcript=transcript,
                agenda_text=agenda_text,
                force_reembed=body.force_reembed,
                k_transcript=body.k_transcript,
                k_meeting=body.k_meeting,
                k_global=body.k_global,
                relative_cutoff=body.relative_similarity_cutoff,
                char_limit=body.char_limit,
                # Agenda-first: skip LLM agenda parsing if frontend supplies them
                agenda_items_override=[a.model_dump() for a in body.agenda_items] if body.agenda_items else None,
                # Hybrid retrieval params
                timeline_stride_seconds=body.timeline_stride_seconds,
                high_confidence_threshold=body.high_confidence_threshold,
                recording_duration=body.recording_duration,
                retrieve_by_timeline=body.retrieve_by_timeline,
                retrieval_mode=body.retrieval_mode,
                max_overlap_chunks=body.max_overlap_chunks,
            ),
        )

    except FileNotFoundError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"[RawMoM] Retrieve pipeline failed for {recording_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Retrieval failed: {str(e)}")

    return result


# ── POST agenda (create/load agenda items only, no LLM if cached) ─────────────

class CreateAgendaRequest(BaseModel):
    force_reparse: bool = False


@router.post("/{recording_id}/agenda")
async def create_agenda_endpoint(
    recording_id: str,
    body: CreateAgendaRequest = CreateAgendaRequest(),
    current_user: dict = Depends(get_current_user),
):
    """
    Generate or load the agenda items for a recording.

    If a cached parsed agenda already exists in the DB and force_reparse=False,
    returns it immediately (no LLM call).  Otherwise runs get_or_create_agenda_items
    which will parse the raw agenda document or generate from the context summary.

    Returns the agenda as an ordered list so the frontend can display and edit
    each item before proceeding to evidence retrieval.
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    transcript = from_json(rec.get("transcript"), [])
    agenda_text: Optional[str] = rec.get("agenda_summary") or None
    if not agenda_text:
        agenda_text = await _load_raw_agenda_text(recording_id, user_id)

    import asyncio
    _loop = asyncio.get_running_loop()

    try:
        agendas, source = await _loop.run_in_executor(
            None,
            lambda: _run_create_agenda(
                recording_id=recording_id,
                user_id=user_id,
                transcript=transcript,
                agenda_text=agenda_text,
                force_reparse=body.force_reparse,
            ),
        )
    except Exception as e:
        logger.error(f"[RawMoM] Agenda creation failed for {recording_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Agenda creation failed: {str(e)}")

    return {"agendas": agendas, "source": source}


# ── POST retrieve-agenda (single-agenda retrieval) ────────────────────────────

class RetrieveAgendaRequest(BaseModel):
    topic: str
    speaker: Optional[str] = None
    agenda_index: int = Field(default=0, ge=0)
    total_agendas: int = Field(default=1, ge=1)
    k_transcript: int = Field(default=8, ge=0, le=30)
    k_meeting: int = Field(default=4, ge=0, le=30)
    k_global: int = Field(default=2, ge=0, le=30)
    relative_similarity_cutoff: float = Field(default=0.01, ge=0.0, le=1.0)
    char_limit: int = Field(default=15000, ge=1000, le=100000)
    force_reembed: bool = False
    timeline_stride_seconds: float = Field(default=60.0, ge=0.0, le=600.0)
    high_confidence_threshold: float = Field(default=0.70, ge=0.0, le=1.0)
    recording_duration: float = Field(default=0.0, ge=0.0)
    retrieve_by_timeline: bool = False
    # Retrieval mode: "agenda_wise" (default) or "chunk_wise"
    # Note: chunk_wise requires all agendas, so per-agenda endpoint ignores this
    # and always uses agenda_wise for single-agenda retrieval.
    retrieval_mode: str = Field(default="agenda_wise")


@router.post("/{recording_id}/retrieve-agenda")
async def retrieve_single_agenda_endpoint(
    recording_id: str,
    body: RetrieveAgendaRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Retrieve evidence for a single agenda item only.

    Embeds transcript/meeting context if not already done (same as /retrieve),
    then runs the hybrid retrieval for the given agenda topic/index.

    Returns a single AgendaResult object (not a list).
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    transcript = from_json(rec.get("transcript"), [])
    if not transcript:
        raise HTTPException(
            status_code=400,
            detail="No transcript available. Run transcription first.",
        )

    import asyncio
    _loop = asyncio.get_running_loop()

    try:
        result = await _loop.run_in_executor(
            None,
            lambda: _run_retrieve_single_agenda(
                recording_id=recording_id,
                user_id=user_id,
                transcript=transcript,
                topic=body.topic,
                speaker=body.speaker,
                agenda_index=body.agenda_index,
                total_agendas=body.total_agendas,
                force_reembed=body.force_reembed,
                k_transcript=body.k_transcript,
                k_meeting=body.k_meeting,
                k_global=body.k_global,
                relative_cutoff=body.relative_similarity_cutoff,
                char_limit=body.char_limit,
                timeline_stride_seconds=body.timeline_stride_seconds,
                high_confidence_threshold=body.high_confidence_threshold,
                recording_duration=body.recording_duration,
                retrieve_by_timeline=body.retrieve_by_timeline,
            ),
        )
    except FileNotFoundError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(
            f"[RawMoM] retrieve-agenda failed for {recording_id} agenda {body.agenda_index}: {e}",
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail=f"Retrieval failed: {str(e)}")

    return result


# ── POST extract-file-text (in-session file text extraction) ─────────────────

from fastapi import UploadFile, File as FastAPIFile
import tempfile
import os as _os


@router.post("/{recording_id}/extract-file-text")
async def extract_file_text_endpoint(
    recording_id: str,
    file: UploadFile = FastAPIFile(...),
    current_user: dict = Depends(get_current_user),
):
    """
    Accept a multipart file upload and extract its plain text.

    No DB storage — the text is returned directly to the caller for
    in-session use as per-agenda context. Supports PDF, DOCX, PPTX, TXT, MD,
    PNG, JPG, XLSX, XLS, CSV (same as doc_extractor).

    Returns { "filename": str, "text": str, "char_count": int }
    """
    # Auth check — ensure the recording belongs to this user
    await _get_recording_or_404(recording_id, current_user["id"])

    filename = file.filename or "upload"
    content = await file.read()

    # Write to temp file for doc_extractor (it works with file paths)
    suffix = _os.path.splitext(filename)[1] or ".bin"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        from services.doc_extractor import extract_text_from_file
        import asyncio
        _loop = asyncio.get_running_loop()
        text = await _loop.run_in_executor(
            None, lambda: extract_text_from_file(tmp_path, filename)
        )
    except Exception as e:
        logger.error(f"[RawMoM] extract-file-text failed for {filename}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(e)}")
    finally:
        try:
            _os.unlink(tmp_path)
        except Exception:
            pass

    text = text.strip()
    return {
        "filename": filename,
        "text": text,
        "char_count": len(text),
    }


# ── POST validate-agenda-similarity ──────────────────────────────────────────

class SimilarityPointInput(BaseModel):
    agenda_idx: int
    disc_idx: int
    point: str


class AgendaSimilarityRequest(BaseModel):
    agenda_topics: List[str]
    points: List[SimilarityPointInput]


@router.post("/{recording_id}/validate-agenda-similarity")
async def validate_agenda_similarity_endpoint(
    recording_id: str,
    body: AgendaSimilarityRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Compute cosine similarity between every discussion point and every agenda topic.

    Uses the same embedding model as the RAG pipeline.  No disk I/O — everything
    is in-memory.  Returns a matrix so the frontend can determine suggested moves.
    """
    await _get_recording_or_404(recording_id, current_user["id"])

    if not body.agenda_topics or not body.points:
        raise HTTPException(status_code=400, detail="agenda_topics and points must not be empty")

    import asyncio
    import numpy as np
    from services.rag_pipeline import _get_embedder

    def _compute() -> dict:
        embedder = _get_embedder()

        # Embed agenda topics
        agenda_embeddings = [
            embedder.encode(topic) for topic in body.agenda_topics
        ]  # list of 1-D numpy arrays

        # Stack into matrix (num_agendas × dim)
        A = np.stack(agenda_embeddings, axis=0)  # shape: (num_agendas, dim)
        # L2-normalise rows for cosine similarity via dot product
        A_norms = np.linalg.norm(A, axis=1, keepdims=True)
        A_norms[A_norms == 0] = 1e-9
        A_norm = A / A_norms

        matrix_rows = []
        for pt in body.points:
            p_emb = embedder.encode(pt.point)  # 1-D array
            p_norm_val = np.linalg.norm(p_emb)
            if p_norm_val == 0:
                p_norm_val = 1e-9
            p_emb_norm = p_emb / p_norm_val

            # Cosine similarities: dot product with each agenda row
            sims = (A_norm @ p_emb_norm).tolist()  # list of floats, one per agenda

            matrix_rows.append({
                "agenda_idx": pt.agenda_idx,
                "disc_idx": pt.disc_idx,
                "point": pt.point,
                "similarities": [round(s, 4) for s in sims],
            })

        return {
            "embeddings_computed": len(body.agenda_topics) + len(body.points),
            "matrix": matrix_rows,
        }

    loop = asyncio.get_running_loop()
    try:
        result = await loop.run_in_executor(None, _compute)
    except Exception as e:
        logger.error(f"[RawMoM] validate-agenda-similarity failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Similarity computation failed: {e}")

    return result


# ── POST generate (from pre-assembled evidence) ───────────────────────────────

@router.post("/{recording_id}/generate")
async def generate_from_evidence_endpoint(
    recording_id: str,
    body: GenerateFromEvidenceRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Generate Raw MoM using pre-assembled evidence provided by the caller.

    This is the second step of the interactive Raw MoM Lab workflow.
    The caller provides the evidence string per agenda item (assembled from
    approved chunks selected in the UI). No FAISS retrieval is performed here.

    The result is saved to recordings.raw_mom and returned.
    """
    user_id = current_user["id"]
    await _get_recording_or_404(recording_id, user_id)

    if not body.agendas:
        raise HTTPException(status_code=400, detail="No agenda items provided.")

    import asyncio
    _loop = asyncio.get_running_loop()

    try:
        raw_mom = await _loop.run_in_executor(
            None,
            lambda: _run_generate_from_evidence(body.agendas),
        )
    except Exception as e:
        logger.error(f"[RawMoM] Generate-from-evidence failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

    # Persist to DB
    raw_mom_json = json.dumps(raw_mom, ensure_ascii=False, default=str)
    async with get_db() as db:
        await db.execute(
            text("UPDATE recordings SET raw_mom = :raw_mom WHERE id = :id AND user_id = :uid"),
            {"raw_mom": raw_mom_json, "id": recording_id, "uid": user_id},
        )
        await db.commit()

    logger.info(f"[RawMoM] Saved evidence-based raw_mom for recording {recording_id}")
    return raw_mom


# ── PUT / POST save raw_mom ───────────────────────────────────────────────────

class SaveRawMomRequest(BaseModel):
    raw_mom: dict


@router.put("/{recording_id}")
@router.post("/{recording_id}/save")
async def save_raw_mom_endpoint(
    recording_id: str,
    body: SaveRawMomRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Save / update the stored raw_mom JSON in DB for a recording.
    """
    user_id = current_user["id"]
    await _get_recording_or_404(recording_id, user_id)

    raw_mom_json = json.dumps(body.raw_mom, ensure_ascii=False, default=str)
    async with get_db() as db:
        await db.execute(
            text("UPDATE recordings SET raw_mom = :raw_mom WHERE id = :id AND user_id = :uid"),
            {"raw_mom": raw_mom_json, "id": recording_id, "uid": user_id},
        )
        await db.commit()

    logger.info(f"[RawMoM] Manually updated raw_mom saved for recording {recording_id}")
    return {"status": "ok", "message": "Raw MoM saved to database"}



# ── POST generate-final-mom (Raw MoM → Final MoM, independent pipeline) ──────

@router.post("/{recording_id}/generate-final-mom")
async def generate_final_mom_from_raw_mom_endpoint(
    recording_id: str,
    req: Optional[GenerateFinalMomRequest] = None,
    current_user: dict = Depends(get_current_user),
):
    """
    Generate a final MoM from the stored Raw MoM JSON.

    This is a COMPLETELY INDEPENDENT pipeline from POST /mom/{id}/generate.
    It does NOT use the transcript. It reads the raw_mom stored by the
    Raw MoM Lab and converts it into a polished final Minutes of Meeting
    using a dedicated AI prompt.

    The result is saved to (or updates) the minutes_of_meeting table and
    returned in the standard MoM schema — immediately viewable in the
    MoM Editor.
    """
    user_id = current_user["id"]
    char_limit = req.char_limit if req else None
    use_max_tokens = req.use_max_tokens if req else False

    # Fetch recording with all needed columns
    async with get_db() as db:
        r = await db.execute(
            text("""
                SELECT id, filename, created_at, duration, speakers_detected, raw_mom
                FROM recordings
                WHERE id = :id AND user_id = :uid
            """),
            {"id": recording_id, "uid": user_id},
        )
        rec = r.mappings().fetchone()

    if not rec:
        raise HTTPException(status_code=404, detail="Recording not found")

    if not rec.get("raw_mom"):
        raise HTTPException(
            status_code=400,
            detail="Raw MoM has not been generated for this recording. "
                   "Use the Raw MoM Lab to generate it first.",
        )

    try:
        raw_mom_data = json.loads(rec["raw_mom"])
    except Exception:
        raise HTTPException(status_code=500, detail="Raw MoM data is corrupted")

    recording_meta = {
        "filename": rec.get("filename", "Meeting Notes"),
        "created_at": rec.get("created_at", ""),
        "duration": rec.get("duration", 0),
        "speakers_detected": from_json(rec.get("speakers_detected"), []),
    }

    import asyncio as _asyncio
    _loop = _asyncio.get_running_loop()

    try:
        mom_data = await _loop.run_in_executor(
            None,
            lambda: _run_generate_mom_from_raw_mom(
                raw_mom_data,
                recording_meta,
                char_limit=char_limit,
                use_max_tokens=use_max_tokens,
            ),
        )
    except Exception as e:
        logger.error(
            f"[RawMoM] generate-final-mom failed for {recording_id}: {e}",
            exc_info=True,
        )
        raise HTTPException(
            status_code=500,
            detail=f"Final MoM generation failed: {str(e)}",
        )

    # Upsert into minutes_of_meeting (same pattern as mom_router.py)
    from database import to_json, dt_to_str
    from datetime import datetime, timezone
    import uuid as _uuid

    now = datetime.now(timezone.utc)
    mom_id = str(_uuid.uuid4())
    initial_version = [{"version": 1, "data": mom_data, "saved_at": dt_to_str(now)}]

    async with get_db() as db:
        r2 = await db.execute(
            text("SELECT id FROM minutes_of_meeting WHERE recording_id = :rid AND user_id = :uid"),
            {"rid": recording_id, "uid": user_id},
        )
        existing = r2.fetchone()

        if existing:
            await db.execute(
                text("""
                    UPDATE minutes_of_meeting SET
                        title = :title, date = :date, duration = :duration,
                        planned_start_time = :planned_start_time,
                        actual_start_time = :actual_start_time,
                        participants = :participants,
                        introduction = :introduction,
                        points_discussed = :points_discussed,
                        action_items = :action_items,
                        conclusion = :conclusion,
                        versions = :versions, is_draft = 0, updated_at = :updated_at
                    WHERE recording_id = :rid AND user_id = :uid
                """),
                {
                    "title": mom_data.get("title", ""),
                    "date": mom_data.get("date", ""),
                    "duration": mom_data.get("duration", 0),
                    "planned_start_time": mom_data.get("planned_start_time", ""),
                    "actual_start_time": mom_data.get("actual_start_time", ""),
                    "participants": to_json(mom_data.get("participants", [])),
                    "introduction": mom_data.get("introduction", ""),
                    "points_discussed": to_json(mom_data.get("points_discussed", [])),
                    "action_items": to_json(mom_data.get("action_items", [])),
                    "conclusion": mom_data.get("conclusion", ""),
                    "versions": to_json(initial_version),
                    "updated_at": dt_to_str(now),
                    "rid": recording_id,
                    "uid": user_id,
                },
            )
        else:
            await db.execute(
                text("""
                    INSERT INTO minutes_of_meeting (
                        id, recording_id, user_id, title, date, duration,
                        planned_start_time, actual_start_time,
                        participants, introduction, points_discussed,
                        action_items, conclusion,
                        versions, is_draft, created_at, updated_at
                    )
                    VALUES (
                        :id, :rid, :uid, :title, :date, :duration,
                        :planned_start_time, :actual_start_time,
                        :participants, :introduction, :points_discussed,
                        :action_items, :conclusion,
                        :versions, 0, :created_at, :updated_at
                    )
                """),
                {
                    "id": mom_id,
                    "rid": recording_id,
                    "uid": user_id,
                    "title": mom_data.get("title", ""),
                    "date": mom_data.get("date", ""),
                    "duration": mom_data.get("duration", 0),
                    "planned_start_time": mom_data.get("planned_start_time", ""),
                    "actual_start_time": mom_data.get("actual_start_time", ""),
                    "participants": to_json(mom_data.get("participants", [])),
                    "introduction": mom_data.get("introduction", ""),
                    "points_discussed": to_json(mom_data.get("points_discussed", [])),
                    "action_items": to_json(mom_data.get("action_items", [])),
                    "conclusion": mom_data.get("conclusion", ""),
                    "versions": to_json(initial_version),
                    "created_at": dt_to_str(now),
                    "updated_at": dt_to_str(now),
                },
            )
        await db.commit()

        # Fetch and return the saved record
        r3 = await db.execute(
            text("SELECT * FROM minutes_of_meeting WHERE recording_id = :rid AND user_id = :uid"),
            {"rid": recording_id, "uid": user_id},
        )
        saved = r3.mappings().fetchone()

    # Re-use the mom_router helper to serialize the row
    from routers.mom_router import _mom_row_to_dict
    return _mom_row_to_dict(saved)


# ── POST generate-full (legacy full pipeline) ─────────────────────────────────


@router.post("/{recording_id}/generate-full")
async def generate_raw_mom_endpoint(
    recording_id: str,
    force_reembed: bool = False,
    current_user: dict = Depends(get_current_user),
):
    """
    Run the full RAG-based Raw MoM extraction pipeline for a recording.

    Pipeline
    --------
    1. Load transcript from DB
    2. Load agenda text from DB (agenda_summary column)
    3. Embed transcript → FAISS (if not already embedded)
    4. Embed meeting context attachments → FAISS (if not already embedded)
    5. Parse agenda → [{topic, speaker}]
    6. For each agenda item: retrieve evidence → LLM extraction
    7. Persist raw_mom JSON to recordings.raw_mom

    Parameters
    ----------
    force_reembed : If true, re-embed transcript and meeting context even if
                    they were previously embedded (useful after transcript updates).

    Returns
    -------
    The generated raw_mom JSON.
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    transcript = from_json(rec.get("transcript"), [])
    if not transcript:
        raise HTTPException(
            status_code=400,
            detail="No transcript available. Run transcription first before generating Raw MoM.",
        )

    # Load agenda text — use the agenda_summary stored after attachment processing.
    # This is the raw/compressed text of agenda attachments, loaded via the
    # existing attachments pipeline (type='agenda').
    agenda_text: Optional[str] = rec.get("agenda_summary") or None

    # If no agenda_summary in recordings, try to load raw agenda attachment text directly
    if not agenda_text:
        agenda_text = await _load_raw_agenda_text(recording_id, user_id)

    logger.info(
        f"[RawMoM] Starting generation: recording={recording_id}, "
        f"transcript_segs={len(transcript)}, has_agenda={bool(agenda_text)}"
    )

    # Run the pipeline in a thread executor (CPU/GPU intensive)
    import asyncio
    _loop = asyncio.get_running_loop()

    try:
        raw_mom = await _loop.run_in_executor(
            None,
            lambda: _run_rag_pipeline(
                recording_id=recording_id,
                user_id=user_id,
                transcript=transcript,
                agenda_text=agenda_text,
                force_reembed_transcript=force_reembed,
                force_reembed_meeting=force_reembed,
            ),
        )
    except FileNotFoundError as e:
        # Embedding model not found
        raise HTTPException(
            status_code=503,
            detail=str(e),
        )
    except Exception as e:
        logger.error(f"[RawMoM] Pipeline failed for {recording_id}: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Raw MoM generation failed: {str(e)}",
        )

    # Persist to DB
    raw_mom_json = json.dumps(raw_mom, ensure_ascii=False, default=str)
    now = dt_to_str(datetime.now(timezone.utc))

    async with get_db() as db:
        await db.execute(
            text("UPDATE recordings SET raw_mom = :raw_mom WHERE id = :id AND user_id = :uid"),
            {"raw_mom": raw_mom_json, "id": recording_id, "uid": user_id},
        )
        await db.commit()

    logger.info(f"[RawMoM] Saved raw_mom for recording {recording_id}")
    return raw_mom


def _run_rag_pipeline(
    recording_id: str,
    user_id: str,
    transcript: list,
    agenda_text: Optional[str],
    force_reembed_transcript: bool,
    force_reembed_meeting: bool,
) -> dict:
    """Synchronous wrapper for the async-hostile RAG pipeline."""
    from services.rag_pipeline import generate_raw_mom
    return generate_raw_mom(
        recording_id=recording_id,
        user_id=user_id,
        transcript=transcript,
        agenda_text=agenda_text,
        force_reembed_transcript=force_reembed_transcript,
        force_reembed_meeting=force_reembed_meeting,
    )


async def _load_raw_agenda_text(recording_id: str, user_id: str) -> Optional[str]:
    """
    Load raw text from agenda attachments directly (fallback when agenda_summary is empty).
    Uses the doc_extractor to get text from the first agenda file.
    """
    try:
        from services.doc_extractor import extract_text_from_file
        async with get_db() as db:
            r = await db.execute(
                text(
                    "SELECT filename, file_path FROM recording_attachments "
                    "WHERE recording_id = :rid AND user_id = :uid AND type = 'agenda' "
                    "ORDER BY created_at ASC LIMIT 1"
                ),
                {"rid": recording_id, "uid": user_id},
            )
            att = r.mappings().fetchone()

        if not att:
            return None

        text_content = extract_text_from_file(att["file_path"], att["filename"])
        return text_content.strip() if text_content else None

    except Exception as e:
        logger.warning(f"[RawMoM] Could not load raw agenda text: {e}")
        return None


# ── GET download/docx ────────────────────────────────────────────────────────

@router.get("/{recording_id}/download/docx")
async def download_raw_mom_docx(
    recording_id: str,
    current_user: dict = Depends(get_current_user),
):
    """
    Generate and download a DOCX version of the Raw MoM as a formatted table.
    Columns: # | Agenda | Speaker | Discussion Points | Actions | Deadline | Owner
    """
    user_id = current_user["id"]
    rec = await _get_recording_or_404(recording_id, user_id)

    if not rec.get("raw_mom"):
        raise HTTPException(status_code=404, detail="Raw MoM has not been generated yet.")

    try:
        raw_mom_data = json.loads(rec["raw_mom"])
    except Exception:
        raise HTTPException(status_code=500, detail="Raw MoM data is corrupted")

    try:
        docx_bytes = _build_docx(raw_mom_data, rec.get("filename", "Recording"))
    except Exception as e:
        logger.error(f"[RawMoM] DOCX generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

    safe_name = rec.get("filename", "raw_mom").replace(" ", "_").rsplit(".", 1)[0]
    filename = f"{safe_name}_raw_mom.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── DELETE raw_mom ────────────────────────────────────────────────────────────

@router.delete("/{recording_id}")
async def delete_raw_mom(
    recording_id: str,
    current_user: dict = Depends(get_current_user),
):
    """
    Clear the raw_mom data for a recording.

    This does NOT delete the FAISS embeddings (transcript/meeting context).
    Use this to re-run the pipeline cleanly.
    """
    user_id = current_user["id"]
    await _get_recording_or_404(recording_id, user_id)

    async with get_db() as db:
        await db.execute(
            text("UPDATE recordings SET raw_mom = NULL WHERE id = :id AND user_id = :uid"),
            {"id": recording_id, "uid": user_id},
        )
        await db.commit()

    return {"status": "cleared", "recording_id": recording_id}


# ── Sync pipeline runners ─────────────────────────────────────────────────────

def _run_create_agenda(
    recording_id: str,
    user_id: str,
    transcript: list,
    agenda_text: Optional[str],
    force_reparse: bool,
) -> tuple:
    """Synchronous agenda generation. Returns (agenda_items, source)."""
    from services.rag_pipeline import (
        get_or_create_agenda_items,
        _load_parsed_agenda,
        _save_parsed_agenda,
    )
    from services.ai_provider import get_provider
    from services.text_embedding_service import unload_text_embedder

    try:
        if not force_reparse:
            cached = _load_parsed_agenda(recording_id)
            if cached:
                logger.info(f"[RawMoM] Returning cached agenda ({len(cached)} items)")
                return cached, "cached"

        if force_reparse:
            # Clear cached agenda so get_or_create re-parses
            _save_parsed_agenda(recording_id, user_id, [])

        agenda_items = get_or_create_agenda_items(
            recording_id=recording_id,
            user_id=user_id,
            transcript=transcript,
            agenda_text=agenda_text,
        )
        source = "parsed" if agenda_text else "generated"
        return agenda_items, source
    finally:
        try:
            from services.ai_provider import QwenProvider
            QwenProvider.unload_model()
        except Exception:
            pass
        try:
            unload_text_embedder()
        except Exception:
            pass



def _run_retrieve_pipeline(
    recording_id: str,
    user_id: str,
    transcript: list,
    agenda_text: Optional[str],
    force_reembed: bool,
    k_transcript: int,
    k_meeting: int,
    k_global: int,
    relative_cutoff: float,
    char_limit: int,
    # Hybrid retrieval params
    agenda_items_override: Optional[list] = None,
    timeline_stride_seconds: float = 60.0,
    high_confidence_threshold: float = 0.70,
    recording_duration: float = 0.0,
    retrieve_by_timeline: bool = False,
    # Retrieval mode
    retrieval_mode: str = "agenda_wise",
    max_overlap_chunks: int = 2,
) -> dict:
    """Synchronous retrieval-only pipeline (no LLM).

    Supports two retrieval modes:
    - agenda_wise  (default): For each agenda, retrieve top-K transcript chunks
                              by semantic similarity to the agenda topic.
    - chunk_wise            : For each transcript chunk, compare it against ALL
                              agendas and assign it to those within the relative
                              similarity threshold of the best-scoring agenda.
    """
    from services.rag_pipeline import (
        embed_transcript, embed_meeting_context,
        retrieve_evidence_raw, retrieve_evidence_chunkwise,
        _transcript_embedded, _meeting_context_embedded,
        _mark_transcript_embedded, _mark_meeting_context_embedded,
        get_or_create_agenda_items,
    )
    from services.text_embedding_service import unload_text_embedder

    try:
        # Step 1: Ensure transcript is embedded
        if transcript and (force_reembed or not _transcript_embedded(recording_id)):
            count = embed_transcript(recording_id, transcript, user_id)
            if count > 0:
                _mark_transcript_embedded(recording_id, user_id)

        # Step 2: Ensure meeting context is embedded
        if force_reembed or not _meeting_context_embedded(recording_id):
            count = embed_meeting_context(recording_id, user_id)
            if count > 0:
                _mark_meeting_context_embedded(recording_id, user_id)

        # Step 3: Get agenda items — use pre-parsed override if provided
        if agenda_items_override:
            agenda_items = agenda_items_override
            logger.info(f"[RawMoM] Using {len(agenda_items)} pre-parsed agenda items from frontend")
        else:
            agenda_items = get_or_create_agenda_items(
                recording_id=recording_id,
                user_id=user_id,
                transcript=transcript,
                agenda_text=agenda_text,
            )

        total_agendas = len(agenda_items)

        # Step 4: Retrieve chunks — branch by retrieval mode
        if retrieval_mode == "chunk_wise":
            logger.info(f"[RawMoM] Running chunk-wise retrieval for {total_agendas} agendas (max_overlap_chunks={max_overlap_chunks})")
            agendas_with_chunks = retrieve_evidence_chunkwise(
                recording_id=recording_id,
                user_id=user_id,
                agenda_items=agenda_items,
                k_global=k_global,
                k_meeting=k_meeting,
                k_transcript=k_transcript,
                relative_cutoff=relative_cutoff,
                char_limit=char_limit,
                recording_duration=recording_duration,
                timeline_stride=timeline_stride_seconds,
                high_confidence_threshold=high_confidence_threshold,
                retrieve_by_timeline=retrieve_by_timeline,
                max_overlap_chunks=max_overlap_chunks,
            )
        else:
            # Default: agenda-wise retrieval (existing behaviour)
            logger.info(f"[RawMoM] Running agenda-wise retrieval for {total_agendas} agendas")
            agendas_with_chunks = []
            for idx, item in enumerate(agenda_items):
                topic = item.get("topic", "")
                if not topic:
                    continue
                chunks = retrieve_evidence_raw(
                    agenda_topic=topic,
                    recording_id=recording_id,
                    user_id=user_id,
                    k_global=k_global,
                    k_meeting=k_meeting,
                    k_transcript=k_transcript,
                    relative_cutoff=relative_cutoff,
                    char_limit=char_limit,
                    agenda_index=idx,
                    total_agendas=total_agendas,
                    recording_duration=recording_duration,
                    timeline_stride=timeline_stride_seconds,
                    high_confidence_threshold=high_confidence_threshold,
                    retrieve_by_timeline=retrieve_by_timeline,
                )
                agendas_with_chunks.append({
                    "topic": topic,
                    "speaker": item.get("speaker"),
                    "is_procedural": chunks["is_procedural"],
                    "transcript_chunks": chunks["transcript"],
                    "meeting_chunks": chunks["meeting"],
                    "global_chunks": chunks["global"],
                })

        return {
            "agendas": agendas_with_chunks,
            "retrieval_params": {
                "k_transcript": k_transcript,
                "k_meeting": k_meeting,
                "k_global": k_global,
                "relative_similarity_cutoff": relative_cutoff,
                "char_limit": char_limit,
                "timeline_stride_seconds": timeline_stride_seconds,
                "high_confidence_threshold": high_confidence_threshold,
                "recording_duration": recording_duration,
                "retrieval_mode": retrieval_mode,
            },
        }
    finally:
        # Unload embedding model to free VRAM
        try:
            unload_text_embedder()
        except Exception as e:
            logger.warning(f"[RawMoM] Failed to unload embedder: {e}")



def _run_retrieve_single_agenda(
    recording_id: str,
    user_id: str,
    transcript: list,
    topic: str,
    speaker: Optional[str],
    agenda_index: int,
    total_agendas: int,
    force_reembed: bool,
    k_transcript: int,
    k_meeting: int,
    k_global: int,
    relative_cutoff: float,
    char_limit: int,
    timeline_stride_seconds: float,
    high_confidence_threshold: float,
    recording_duration: float,
    retrieve_by_timeline: bool = False,
) -> dict:
    """Synchronous retrieval for a single agenda item. Embeds if needed."""
    from services.rag_pipeline import (
        embed_transcript, embed_meeting_context,
        retrieve_evidence_raw, _transcript_embedded, _meeting_context_embedded,
        _mark_transcript_embedded, _mark_meeting_context_embedded,
    )
    from services.text_embedding_service import unload_text_embedder

    try:
        # Ensure transcript is embedded
        if transcript and (force_reembed or not _transcript_embedded(recording_id)):
            count = embed_transcript(recording_id, transcript, user_id)
            if count > 0:
                _mark_transcript_embedded(recording_id, user_id)

        # Ensure meeting context is embedded
        if force_reembed or not _meeting_context_embedded(recording_id):
            count = embed_meeting_context(recording_id, user_id)
            if count > 0:
                _mark_meeting_context_embedded(recording_id, user_id)

        chunks = retrieve_evidence_raw(
            agenda_topic=topic,
            recording_id=recording_id,
            user_id=user_id,
            k_global=k_global,
            k_meeting=k_meeting,
            k_transcript=k_transcript,
            relative_cutoff=relative_cutoff,
            char_limit=char_limit,
            agenda_index=agenda_index,
            total_agendas=total_agendas,
            recording_duration=recording_duration,
            timeline_stride=timeline_stride_seconds,
            high_confidence_threshold=high_confidence_threshold,
            retrieve_by_timeline=retrieve_by_timeline,
        )

        return {
            "topic": topic,
            "speaker": speaker,
            "is_procedural": chunks["is_procedural"],
            "transcript_chunks": chunks["transcript"],
            "meeting_chunks": chunks["meeting"],
            "global_chunks": chunks["global"],
        }
    finally:
        try:
            unload_text_embedder()
        except Exception as e:
            logger.warning(f"[RawMoM] Failed to unload embedder: {e}")


def _run_generate_from_evidence(agendas: list) -> dict:
    """Synchronous LLM generation using pre-assembled evidence (no retrieval)."""
    import gc
    import torch
    from services.ai_provider import get_provider, QwenProvider

    provider = get_provider()
    processed_agendas = []

    try:
        for item in agendas:
            topic = item.topic
            speaker = item.speaker
            evidence = item.evidence

            # Append per-agenda manual context if provided
            manual_ctx = (item.manual_context or "").strip()
            if manual_ctx:
                evidence = evidence + "\n\n=== AGENDA-SPECIFIC CONTEXT ===\n" + manual_ctx

            try:
                result = provider.extract_raw_mom_for_agenda(
                    agenda_topic=topic,
                    agenda_speaker=speaker,
                    evidence=evidence,
                )
                processed_agendas.append(result)
                logger.info(
                    f"[RawMoM] '{topic[:40]}': "
                    f"{len(result.get('discussion', []))} discussion entries"
                )
            except Exception as e:
                logger.error(f"[RawMoM] Extraction failed for '{topic}': {e}")
                processed_agendas.append({
                    "agenda_topic": topic,
                    "agenda_speaker": speaker,
                    "discussion": [],
                })
            finally:
                gc.collect()
                try:
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
                except Exception:
                    pass

        return {"meeting": {"agendas": processed_agendas}}
    finally:
        try:
            QwenProvider.unload_model()
        except Exception as e:
            logger.warning(f"[RawMoM] Failed to unload LLM: {e}")


def _run_generate_mom_from_raw_mom(
    raw_mom_data: dict,
    recording_meta: dict,
    char_limit: Optional[int] = None,
    use_max_tokens: Optional[bool] = False,
) -> dict:
    """Synchronous wrapper for the Raw MoM → Final MoM conversion pipeline."""
    from services.ai_provider import QwenProvider
    from services.llm import generate_mom_from_raw_mom
    try:
        return generate_mom_from_raw_mom(
            raw_mom_data,
            recording_meta,
            char_limit=char_limit,
            use_max_tokens=use_max_tokens,
        )
    finally:
        try:
            QwenProvider.unload_model()
        except Exception as e:
            logger.warning(f"[RawMoM] Failed to unload LLM after generate-final-mom: {e}")


def _build_docx(raw_mom_data: dict, recording_name: str) -> bytes:
    """Build a DOCX file from raw_mom JSON and return bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )

    doc = Document()

    # Title
    title = doc.add_heading(f"Raw Minutes of Meeting", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Recording: {recording_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    agendas = raw_mom_data.get("meeting", {}).get("agendas", [])
    if not agendas:
        doc.add_paragraph("No agenda items found.")
    else:
        # Table headers
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = ["Agenda ID", "Agenda", "Point", "Action"]
        hdr = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr[i].text = header
            hdr[i].paragraphs[0].runs[0].bold = True

        agenda_id = 1

        for agenda in agendas:
            topic = agenda.get("agenda_topic", "")
            default_speaker = agenda.get("agenda_speaker")
            discussion = agenda.get("discussion", [])

            if not discussion:
                row = table.add_row().cells
                row[0].text = str(agenda_id)
                row[1].text = topic
                row[2].text = ""
                row[3].text = "For Information"
            else:
                first_row = True

                for entry in discussion:
                    row = table.add_row().cells

                    # Show Agenda ID and Agenda only on the first row
                    if first_row:
                        row[0].text = str(agenda_id)
                        row[1].text = topic
                        first_row = False

                    row[2].text = str(entry.get("point") or "")

                    speaker = (
                        entry.get("speaker")
                        or default_speaker
                        or "For Information"
                    )
                    row[3].text = str(speaker)

            agenda_id += 1
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
